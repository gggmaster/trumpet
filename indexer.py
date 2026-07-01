import argparse
import hashlib
import html
import json
import os
import re
import sqlite3
import sys
import time
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


SOURCE_ROOT = Path(
    os.environ.get(
        "INVESTMENT_SOURCE",
        r"C:\Users\ll_ga\OneDrive\00\investment",
    )
)
DB_PATH = Path(os.environ.get("INVESTMENT_DB", Path(__file__).with_name("investment_index.db")))
STATUS_PATH = Path(__file__).with_name("index_status.json")
SKIP_PDF = False

SUPPORTED = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".html", ".htm"}
IDENTITY_PATTERN = re.compile(
    r"(passport|driver.?licen[cs]e|identity.?card|medicare|birth.?certificate)",
    re.IGNORECASE,
)
SPACE_RE = re.compile(r"\s+")
TAG_RE = re.compile(r"<[^>]+>")


def clean_text(value):
    if value is None:
        return ""
    return SPACE_RE.sub(" ", str(value).replace("\x00", " ")).strip()


def chunks(text, size=1600, overlap=220):
    text = clean_text(text)
    if not text:
        return
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        if end < len(text):
            split = text.rfind(" ", start + size // 2, end)
            if split > start:
                end = split
        piece = text[start:end].strip()
        if piece:
            yield piece
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)


def extract_pdf(path):
    reader = PdfReader(str(path))
    for page_no, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        for part_no, text_part in enumerate(chunks(text), 1):
            yield {
                "locator": f"page {page_no}",
                "part": part_no,
                "text": text_part,
            }


def extract_docx(path):
    doc = Document(str(path))
    pieces = []
    for paragraph in doc.paragraphs:
        if clean_text(paragraph.text):
            pieces.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            pieces.append(" | ".join(clean_text(cell.text) for cell in row.cells))
    for part_no, text_part in enumerate(chunks("\n".join(pieces)), 1):
        yield {"locator": "document", "part": part_no, "text": text_part}


def extract_workbook(path):
    book = load_workbook(str(path), read_only=True, data_only=True)
    try:
        for sheet in book.worksheets:
            buffer = []
            part_no = 0
            for row in sheet.iter_rows(values_only=True):
                values = [clean_text(value) for value in row]
                if not any(values):
                    continue
                buffer.append(" | ".join(values))
                if sum(len(line) for line in buffer) >= 1400:
                    part_no += 1
                    yield {
                        "locator": f"sheet {sheet.title}",
                        "part": part_no,
                        "text": clean_text("\n".join(buffer)),
                    }
                    buffer = []
            if buffer:
                part_no += 1
                yield {
                    "locator": f"sheet {sheet.title}",
                    "part": part_no,
                    "text": clean_text("\n".join(buffer)),
                }
    finally:
        book.close()


def extract_text_file(path):
    raw = path.read_text(encoding="utf-8", errors="ignore")
    if path.suffix.lower() in {".html", ".htm"}:
        raw = html.unescape(TAG_RE.sub(" ", raw))
    for part_no, text_part in enumerate(chunks(raw), 1):
        yield {"locator": "document", "part": part_no, "text": text_part}


def extract(path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        yield from extract_pdf(path)
    elif suffix == ".docx":
        yield from extract_docx(path)
    elif suffix in {".xlsx", ".xlsm"}:
        yield from extract_workbook(path)
    else:
        yield from extract_text_file(path)


def init_database(connection):
    connection.executescript(
        """
        PRAGMA journal_mode=WAL;
        DROP TABLE IF EXISTS chunks_fts;
        DROP TABLE IF EXISTS chunks;
        DROP TABLE IF EXISTS files;

        CREATE TABLE files (
            id INTEGER PRIMARY KEY,
            path TEXT NOT NULL UNIQUE,
            relative_path TEXT NOT NULL,
            extension TEXT,
            size INTEGER,
            modified REAL,
            sha1 TEXT,
            status TEXT NOT NULL,
            error TEXT
        );

        CREATE TABLE chunks (
            id INTEGER PRIMARY KEY,
            file_id INTEGER NOT NULL,
            locator TEXT,
            part INTEGER,
            text TEXT NOT NULL,
            FOREIGN KEY(file_id) REFERENCES files(id)
        );

        CREATE VIRTUAL TABLE chunks_fts USING fts5(
            text,
            relative_path,
            tokenize='unicode61 remove_diacritics 2'
        );
        """
    )


def file_sha1(path):
    digest = hashlib.sha1()
    with path.open("rb") as stream:
        while True:
            block = stream.read(1024 * 1024)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def write_status(stats, started, complete=False):
    payload = dict(stats)
    payload.update(
        {
            "source_root": str(SOURCE_ROOT),
            "database": str(DB_PATH),
            "started_at": started,
            "updated_at": time.time(),
            "complete": complete,
        }
    )
    STATUS_PATH.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def build_index():
    if not SOURCE_ROOT.exists():
        raise SystemExit(f"Source folder not found: {SOURCE_ROOT}")

    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    connection = sqlite3.connect(DB_PATH)
    init_database(connection)
    started = time.time()
    stats = {
        "total_files": 0,
        "indexed_files": 0,
        "indexed_chunks": 0,
        "skipped_files": 0,
        "identity_files_skipped": 0,
        "failed_files": 0,
        "offline_or_unavailable": 0,
        "unsupported_by_extension": {},
        "errors": [],
    }

    all_files = sorted((p for p in SOURCE_ROOT.rglob("*") if p.is_file()), key=str)
    stats["total_files"] = len(all_files)
    write_status(stats, started)

    for number, path in enumerate(all_files, 1):
        relative = str(path.relative_to(SOURCE_ROOT))
        suffix = path.suffix.lower()
        status = "pending"
        error = None
        file_id = None

        try:
            stat = path.stat()
            if SKIP_PDF and suffix == ".pdf":
                status = "skipped_pdf"
                stats["skipped_files"] += 1
            elif IDENTITY_PATTERN.search(relative):
                status = "skipped_identity"
                stats["identity_files_skipped"] += 1
                stats["skipped_files"] += 1
            elif suffix not in SUPPORTED:
                status = "unsupported"
                stats["skipped_files"] += 1
                label = suffix or "(no extension)"
                stats["unsupported_by_extension"][label] = (
                    stats["unsupported_by_extension"].get(label, 0) + 1
                )

            cursor = connection.execute(
                """
                INSERT INTO files(path, relative_path, extension, size, modified, sha1, status, error)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    str(path),
                    relative,
                    suffix,
                    stat.st_size,
                    stat.st_mtime,
                    None,
                    status,
                    error,
                ),
            )
            file_id = cursor.lastrowid

            if status == "pending":
                chunk_count = 0
                for item in extract(path):
                    text = clean_text(item["text"])
                    if len(text) < 20:
                        continue
                    chunk_cursor = connection.execute(
                        "INSERT INTO chunks(file_id, locator, part, text) VALUES (?, ?, ?, ?)",
                        (file_id, item["locator"], item["part"], text),
                    )
                    connection.execute(
                        "INSERT INTO chunks_fts(rowid, text, relative_path) VALUES (?, ?, ?)",
                        (chunk_cursor.lastrowid, text, relative),
                    )
                    chunk_count += 1
                status = "indexed" if chunk_count else "no_text"
                connection.execute(
                    "UPDATE files SET status=?, sha1=? WHERE id=?",
                    (status, file_sha1(path), file_id),
                )
                if status == "indexed":
                    stats["indexed_files"] += 1
                    stats["indexed_chunks"] += chunk_count
                else:
                    stats["skipped_files"] += 1

        except Exception as exc:
            message = f"{type(exc).__name__}: {exc}"
            lowered = message.lower()
            status = "failed"
            stats["failed_files"] += 1
            if any(term in lowered for term in ("offline", "cloud", "unavailable", "access")):
                stats["offline_or_unavailable"] += 1
            if len(stats["errors"]) < 100:
                stats["errors"].append({"file": relative, "error": message[:500]})
            if file_id is None:
                connection.execute(
                    """
                    INSERT OR IGNORE INTO files
                    (path, relative_path, extension, size, modified, sha1, status, error)
                    VALUES (?, ?, ?, 0, 0, NULL, 'failed', ?)
                    """,
                    (str(path), relative, suffix, message[:1000]),
                )
            else:
                connection.execute(
                    "UPDATE files SET status='failed', error=? WHERE id=?",
                    (message[:1000], file_id),
                )

        if number % 10 == 0 or number == len(all_files):
            connection.commit()
            write_status(stats, started)
            print(
                f"[{number}/{len(all_files)}] indexed={stats['indexed_files']} "
                f"chunks={stats['indexed_chunks']} failed={stats['failed_files']}",
                flush=True,
            )

    connection.commit()
    connection.execute("INSERT INTO chunks_fts(chunks_fts) VALUES('optimize')")
    connection.commit()
    connection.close()
    write_status(stats, started, complete=True)
    print(json.dumps(stats, indent=2))


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Index the OneDrive investment folder.")
    parser.add_argument("--source", help="Override source folder")
    parser.add_argument("--db", help="Override database path")
    parser.add_argument("--skip-pdf", action="store_true", help="Skip PDF files")
    args = parser.parse_args()
    if args.source:
        SOURCE_ROOT = Path(args.source)
    if args.db:
        DB_PATH = Path(args.db)
    SKIP_PDF = args.skip_pdf
    build_index()
